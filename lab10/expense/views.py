import io  #для работы с буфером
from docx import Document  #для генерации Word-файлов
from docx.shared import Pt #для шрифта Word-файлов
from django.shortcuts import render, redirect
from django.http import HttpResponse
from .models import *
from .forms import *
from django.db.models import Sum

#главная страница
def index(request):
    return render(request, 'index.html')


#таблица издержек
def expense_table(request):
    expenses = Expense.objects.all().order_by('-date')
    return render(request, 'table.html', {'expenses': expenses})


#добавление издержки в таблицу
def add_expense(request):
    if request.method == 'POST':
        form = ExpenseForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('expense_table')
    else:
        form = ExpenseForm()
    return render(request, 'add_expense.html', {'form': form})


# создание отчета/сохранение в .docx
def create_report(request):
    form = ReportForm(request.GET or None)
    report_data = None
    total_sum = 0

    if request.GET and form.is_valid():
        title = form.cleaned_data['title']
        exp_type = form.cleaned_data['type']
        start = form.cleaned_data['start_date']
        end = form.cleaned_data['end_date']

        #упорядочить по дате
        report_data = Expense.objects.filter(
            type=exp_type,
            date__range=[start, end]
        ).order_by('date')

        #общая сумма
        total_sum = report_data.aggregate(Sum('amount'))['amount__sum'] or 0

        #скачивание файла в .docx
        if 'download' in request.GET:
            type_display = dict(Expense.TYPE_CHOICES).get(exp_type)
            #создание документа
            doc = Document()
            heading = doc.add_heading(f"{title}", level=1) #заголовок
            heading.alignment = 1  #выравнивание по центру

            #новый абзац
            p_info = doc.add_paragraph()
            p_info.add_run(f"Тип издержек: ").bold = True
            p_info.add_run(f"{type_display}\n")
            p_info.add_run(f"Период: ").bold = True
            p_info.add_run(f"с {start.strftime('%d.%m.%Y')} по {end.strftime('%d.%m.%Y')}\n")

            #таблица (1 строка для заголовков, 2 колонки)
            table = doc.add_table(rows=1, cols=2)
            #шапка таблицы
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Дата учёта'
            hdr_cells[1].text = 'Сумма (руб.)'
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True #текст в шапке жирным шрифтом

            #заполнение таблицы
            for item in report_data:
                row_cells = table.add_row().cells
                row_cells[0].text = item.date.strftime('%d.%m.%Y')
                row_cells[1].text = f"{item.amount:,.2f}".replace(",", " ")

            doc.add_paragraph()  #новый абзац
            p_total = doc.add_paragraph()
            run_total = p_total.add_run(f"ОБЩАЯ СУММА: {total_sum:,.2f} руб.")
            run_total.font.size = Pt(14)
            run_total.bold = True #пишем итоговую сумму жирным шрифтом

            #сохраняем документ в буфер памяти
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="report_{start}_{end}.docx"'
            return response

    return render(request, 'report.html', {
        'form': form,
        'report_data': report_data,
        'total_sum': total_sum
    })